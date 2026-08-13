import docx
from typing import List
from .models import TextBlock, RunSpan

class DocumentReader:
    """Reads DOCX files into structured text blocks while tracking XML run boundaries."""

    def __init__(self, filepath: str):
        self.filepath = filepath
        self.doc = docx.Document(filepath)

    def extract_blocks(()) -> List[TextBlock]:
        blocks: List[TextBlock] = []

        # 1. Main Document Paragraphs
        for idx, p in enumerate(self.doc.paragraphs):
            block = self._build_block(f"p_{idx}", "paragraph", p)
            if block.text.strip():
                blocks.append(block)

        # 2. Document Tables
        for t_idx, table in enumerate(self.doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, p in enumerate(cell.paragraphs):
                        block_id = f"t_{t_idx}_r_{r_idx}_c_{c_idx}_p_{p_idx}"
                        block = self._build_block(block_id, "cell", p)
                        block.row_index = r_idx
                        block.col_index = c_idx
                        if block.text.strip():
                            blocks.append(block)

        # 3. Headers and Footers
        for s_idx, section in enumerate(self.doc.sections):
            if section.header:
                for p_idx, p in enumerate(section.header.paragraphs):
                    block = self._build_block(f"h_{s_idx}_p_{p_idx}", "header", p)
                    block.section_index = s_idx
                    if block.text.strip():
                        blocks.append(block)
            if section.footer:
                for p_idx, p in enumerate(section.footer.paragraphs):
                    block = self._build_block(f"f_{s_idx}_p_{p_idx}", "footer", p)
                    block.section_index = s_idx
                    if block.text.strip():
                        blocks.append(block)

        return blocks

    def _build_block(self, block_id: str, block_type: str, paragraph: docx.text.paragraph.Paragraph) -> TextBlock:
        full_text = ""
        runs: List[RunSpan] = []

        for r_idx, run in enumerate(paragraph.runs):
            run_text = run.text
            start = len(full_text)
            full_text += run_text
            end = len(full_text)

            runs.append(
                RunSpan(
                    run_index=r_idx,
                    start_char=start,
                    end_char=end,
                    text=run_text,
                    run_obj=run,
                )
            )

        return TextBlock(
            block_id=block_id,
            block_type=block_type,
            text=full_text,
            runs=runs,
            raw_element=paragraph,
        )
